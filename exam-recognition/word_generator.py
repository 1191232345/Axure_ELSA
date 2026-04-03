from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

class WordGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def add_title(self, title_text):
        title = self.doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(18)
            run.font.bold = True
    
    def add_section(self, section_data):
        heading = self.doc.add_heading(section_data['title'], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True
        
        for question in section_data.get('questions', []):
            self.add_question(question, section_data.get('type'))
        
        self.doc.add_paragraph()
    
    def add_question(self, question_data, question_type='unknown'):
        paragraph = self.doc.add_paragraph()
        
        question_text = f"{question_data['number']}. {question_data['content']}"
        if question_data.get('score'):
            question_text += f" ({question_data['score']}分)"
        
        run = paragraph.add_run(question_text)
        run.font.size = Pt(12)
        run.font.bold = True
        
        if question_type == 'choice' and question_data.get('options'):
            for option in question_data['options']:
                option_paragraph = self.doc.add_paragraph()
                option_run = option_paragraph.add_run(f"{option['letter']}. {option['content']}")
                option_run.font.size = Pt(11)
                option_paragraph.paragraph_format.left_indent = Inches(0.3)
        
        if question_type in ['fill_blank', 'short_answer', 'calculation', 'application']:
            answer_paragraph = self.doc.add_paragraph()
            answer_paragraph.add_run('_' * 50)
            answer_paragraph.paragraph_format.left_indent = Inches(0.3)
        
        self.doc.add_paragraph()
    
    def generate_from_parsed_data(self, parsed_data, title='试卷'):
        self.add_title(title)
        
        for section in parsed_data.get('sections', []):
            self.add_section(section)
        
        return self.doc
    
    def generate_from_text(self, text, title='试卷'):
        from question_parser import QuestionParser
        
        parser = QuestionParser()
        parsed_data = parser.parse(text)
        formatted_data = parser.format_for_word(parsed_data)
        
        self.add_title(title)
        
        for section in formatted_data:
            self.add_section(section)
        
        return self.doc
    
    def save(self, filename):
        try:
            self.doc.save(filename)
            return True
        except Exception as e:
            print(f"Error saving document: {e}")
            return False
    
    def get_document_bytes(self):
        from io import BytesIO
        
        file_stream = BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
